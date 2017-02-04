from docx import Document

#Employee information stored as list of dictionaries
employee_data = [
    {'id': 123, 'name': 'John Sally', 'department': 'Operations', 'isDue': True},
    {'id': 245, 'name': 'Robert Langford', 'department': 'Software', 'isDue': False},
]

#Agenda of new hire orientation program for different business units
agenda = {
    "Operations": ["SAP Overview", "Inventory Management"],
    "Software": ["C/C++ Overview", "Computer Architecture"],
    "Hardware": ["Computer Aided Tools", "Hardware Design"]
}

#Method to generate a Word document as invite
#based on business unit of employee
def generate_document(employee_data, agenda):
    document = Document()
    for emp in employee_data:
        if emp['isDue']:
            name = emp['name']
            document.add_heading('Your New Hire Orientation\n', level=1)
            document.add_paragraph('Dear %s,' % name)
            document.add_paragraph('Welcome to Google Inc. You have been selected for our new hire orientation.')
            document.add_paragraph('Based on your department you will go through below sessions:')
            department = emp['department']
            for session in agenda[department]:
                document.add_paragraph(
                    session , style='ListBullet'
                )
            document.add_paragraph('Thanks,\n HR Manager')
            document.save('orientation_%s.docx' % emp['id'])

generate_document(employee_data, agenda)
