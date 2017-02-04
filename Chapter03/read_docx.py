import docx

#Get the document object
doc = docx.Document('WExercise.docx')
print "Document Object:", doc

#Print the title text of the Word document
print "Title of the document:"
print doc.paragraphs[0].text

#Go through the cells of the tables added to the Word doc
table = doc.tables[0]
print "Column 1:"
for i in range(len(table.rows)):
    print table.rows[i].cells[0].paragraphs[0].text
print "Column 2:"
for i in range(len(table.rows)):
    print table.rows[i].cells[1].paragraphs[0].text
print "Column 3:"
for i in range(len(table.rows)):
    print table.rows[i].cells[2].paragraphs[0].text

#Read attributes like author, created date for the Word doc
print "Attributes of the document"
print "Author:", doc.core_properties.author
print "Date Created:", doc.core_properties.created
print "Document Revision:", doc.core_properties.revision
